from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "shime" / "deliverables" / "SHIME_20260808_クライアント準備依頼書.docx"

INK = "2E2B31"
MUTED = "66616A"
ACCENT = "C54967"
ACCENT_DARK = "96364D"
PALE = "FCEEF2"
GRAY = "F3F2F4"
WHITE = "FFFFFF"
BORDER = "D8D3D6"
CAUTION = "9E1E1E"
FONT_LATIN = "Calibri"
FONT_JP = "Yu Gothic"
CONTENT_DXA = 9360


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_JP)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_paragraph(doc, text="", size=11, bold=False, color=INK, after=6, before=0, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        set_font(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_checkbox(doc, text, checked=False, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(("☒" if checked else "☐") + " " + text), size=10.5)
    return p


def add_response_box(doc, prompt="回答／添付ファイル名：", lines=2):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, WHITE)
    set_cell_border(cell, BORDER, 6)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(prompt), size=9.5, bold=True, color=MUTED)
    for _ in range(lines):
        q = cell.add_paragraph("　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　")
        q.paragraph_format.space_after = Pt(1)
        set_font(q.runs[0], size=10.5, color=INK)
    add_paragraph(doc, after=5)


def add_question(doc, number, title, guidance, response_lines=2):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(f"{number}. {title}"), size=13, bold=True, color=ACCENT_DARK)
    add_paragraph(doc, guidance, size=10.5, color=INK, after=5)
    add_response_box(doc, lines=response_lines)


def add_section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=16, bold=True, color=ACCENT_DARK)
    return p


def add_note_box(doc, title, body, caution=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF3F3" if caution else PALE)
    set_cell_border(cell, CAUTION if caution else ACCENT, 7)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(title), size=10.5, bold=True, color=CAUTION if caution else ACCENT_DARK)
    q = cell.add_paragraph(body)
    q.paragraph_format.space_after = Pt(0)
    q.paragraph_format.line_spacing = 1.10
    for run in q.runs:
        set_font(run, size=10, color=INK)
    add_paragraph(doc, after=5)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    for i, (label, value) in enumerate(rows):
        if i > 0:
            table.add_row()
            set_table_geometry(table, [2700, 6660])
        left, right = table.rows[i].cells
        set_cell_shading(left, GRAY)
        set_cell_shading(right, WHITE)
        set_cell_border(left)
        set_cell_border(right)
        left.paragraphs[0].paragraph_format.space_after = Pt(0)
        right.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_font(left.paragraphs[0].add_run(label), size=10, bold=True)
        set_font(right.paragraphs[0].add_run(value), size=10)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT_DARK, 16, 8),
        ("Heading 2", 13, ACCENT_DARK, 12, 6),
        ("Heading 3", 12, INK, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = FONT_LATIN
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.10


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(hp.add_run("SHIME® | 2026年8月8日イベント"), size=8.5, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)

    add_paragraph(doc, "CLIENT INFORMATION REQUEST", size=9.5, bold=True, color=ACCENT, after=3)
    add_paragraph(doc, "SHIME® 本番準備情報・資料 提出依頼書", size=25, bold=True, color=INK, after=6)
    add_paragraph(doc, "2026年8月8日のイベント本番設定および運用準備のための確認書", size=13, color=MUTED, after=16)

    add_key_value_table(doc, [
        ("ご提出元", "イベント主催者・運営責任者"),
        ("対象イベント", "2026年8月8日開催予定 SHIMEイベント"),
        ("回答期限", "　　　　年　　　月　　　日"),
        ("回答責任者", "氏名：　　　　　　　　　　　　役職：　　　　　　　　　"),
    ])
    add_paragraph(doc, after=8)
    add_note_box(
        doc,
        "ご提出の目的",
        "本書は、申込受付、LINE本人連携、席配置、希望入力、結果通知を本番環境で安全に運用するための確定情報を収集するものです。未確定欄は空欄にせず「未確定」と記載し、確定予定日と担当者を添えてください。",
    )
    add_note_box(
        doc,
        "重要：個人情報を記載しないでください",
        "本書には本番参加者の氏名、電話番号、メールアドレス、生年月日、LINE情報を記載しないでください。参加者名簿は本書とは別の承認済み手順で取り扱います。",
        caution=True,
    )

    add_section_heading(doc, "提出物チェック")
    for item in (
        "本書の必須14項目への回答",
        "イベント利用規約の最終本文（WordまたはPDF）",
        "プライバシーポリシーの最終本文（WordまたはPDF）",
        "会場レイアウト図と50名分のテーブル・席構成",
        "ロゴ・イベント表記などのブランド素材（使用する場合）",
        "感情カード、Dream固定文、席案内5問への承認または修正指示",
        "当日運用責任者・緊急連絡先一覧（個人情報管理ルールに従い別送可）",
    ):
        add_checkbox(doc, item)

    # Cover content already fills the first page; allow normal pagination to
    # avoid producing a blank page when Word moves the break paragraph.
    add_section_heading(doc, "A. 登録済み設定案の最終確認")
    add_paragraph(doc, "以下は現在の設定案です。「承認」または「変更」を選び、変更がある場合は内容を記入してください。", size=10.5)
    settings = [
        ("開催日", "2026年8月8日"),
        ("開始時刻", "14:00（日本時間）"),
        ("定員", "50名"),
        ("Dream登録", "必須。ただし非公開を選択可能"),
        ("希望入力", "最大2名・順位なし"),
        ("複数成立", "許可しない"),
        ("成立後の連絡", "運営が仲介"),
        ("参加者番号", "区分A：A01〜／区分B：B01〜、2桁"),
        ("AI障害時", "固定文候補へ切り替えて運用を継続"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2400, 5100, 1860])
    headers = ("設定項目", "現在の設定案", "確認")
    mark_header_row(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, ACCENT_DARK)
        set_cell_border(cell, ACCENT_DARK)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cell.paragraphs[0].add_run(value), size=9.5, bold=True, color=WHITE)
    for label, value in settings:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_border(cell)
        set_cell_shading(cells[0], GRAY)
        set_cell_shading(cells[1], WHITE)
        set_cell_shading(cells[2], WHITE)
        set_font(cells[0].paragraphs[0].add_run(label), size=9.5, bold=True)
        set_font(cells[1].paragraphs[0].add_run(value), size=9.5)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cells[2].paragraphs[0].add_run("☐承認\n☐変更"), size=9)
    set_table_geometry(table, [2400, 5100, 1860])
    add_response_box(doc, "変更内容：", lines=3)

    add_section_heading(doc, "B. 本番設定に必要な必須14項目")
    add_note_box(doc, "日時の記入方法", "日時はすべて日本時間（JST）で、年・月・日・時刻まで記載してください。")
    add_question(doc, 1, "正式イベント名", "申込画面、LINE画面、SHIME PASS、管理画面、通知に表示する正式名称です。")
    add_question(doc, 2, "イベント終了日時", "終了予定時刻を記載してください。延長の可能性がある場合は、最終終了時刻も併記してください。")
    add_question(doc, 3, "会場の正式名称", "施設名、会場名、部屋名を正式表記で記載してください。")
    add_question(doc, 4, "会場住所", "郵便番号、都道府県、市区町村、番地、建物名、階・部屋番号まで記載してください。", 3)
    add_question(doc, 5, "申込受付開始日時", "利用者向け申込フォームを公開する日時です。")
    add_question(doc, 6, "申込受付終了日時", "申込フォームを締め切る日時です。締切後の個別受付方針があれば記載してください。", 3)
    add_question(doc, 7, "希望入力開始日時", "参加者が会話相手への希望を入力できるようになる日時です。")
    add_question(doc, 8, "希望入力締切日時", "希望入力を締め切る日時です。結果集計・責任者確認の時間を確保してください。")
    add_question(doc, 9, "参加区分A・Bの正式名称", "利用者に表示する区分名を記載してください。参加者番号の接頭辞A／Bを変更する場合も併記してください。", 3)
    add_question(doc, 10, "席替え回数", "初回着席を除き、イベント中に実施する席替え回数を記載してください。運用上の1回の時間も分かれば併記してください。", 3)
    add_question(doc, 11, "データ保存期間", "イベント終了後、申込・回答・受付・希望・結果等を保存する日数を記載してください。法令、契約、問い合わせ対応期間を踏まえて決定してください。", 3)
    add_question(doc, 12, "イベント利用規約の版番号", "最終本文を別ファイルで提出し、版番号を記載してください。例：2026-08-08-v1", 3)
    add_question(doc, 13, "プライバシーポリシーの版番号", "最終本文を別ファイルで提出し、版番号を記載してください。例：2026-08-08-v1", 3)
    add_question(doc, 14, "テーブル・席構成", "50名分について、テーブル数、各卓の定員、席番号、使用しない席、予備席、配慮席を記載し、会場レイアウト図を添付してください。", 5)

    doc.add_page_break()
    add_section_heading(doc, "C. 規約・法務資料の提出要件")
    add_note_box(
        doc,
        "法務確認について",
        "SHIME側では文書の版管理と同意記録を実装しますが、規約・プライバシーポリシー本文の法的妥当性を保証するものではありません。主催者または主催者が指定する専門家による最終確認をお願いします。",
        caution=True,
    )

    add_section_heading(doc, "C-1. イベント利用規約")
    add_paragraph(doc, "WordまたはPDFで最終本文をご提出ください。少なくとも以下の項目を含めるか、該当しない理由を明記してください。", size=10.5)
    for item in (
        "主催者・サービス提供者の正式名称と連絡先",
        "参加資格、申込条件、本人確認",
        "参加料金、支払方法、キャンセル・返金条件",
        "禁止事項、迷惑行為、参加拒否・退場条件",
        "イベント内容、日時、会場の変更・中止",
        "席配置、希望入力、成立結果に関する説明",
        "成立や交際等を保証しない旨",
        "安全管理、免責、損害賠償の範囲",
        "写真・動画撮影を行う場合の条件",
        "準拠法、管轄、問い合わせ窓口",
        "施行日・版番号",
    ):
        add_checkbox(doc, item)
    add_response_box(doc, "添付ファイル名／法務確認者／確認日：", lines=3)

    add_section_heading(doc, "C-2. プライバシーポリシー")
    add_paragraph(doc, "WordまたはPDFで最終本文をご提出ください。少なくとも以下の項目を含めるか、該当しない理由を明記してください。", size=10.5)
    for item in (
        "個人情報取扱事業者の正式名称、所在地、連絡先",
        "取得する情報（申込情報、LINE識別情報、Dream、質問回答、希望、受付・席・結果等）",
        "利用目的",
        "要配慮情報を取得する場合の取扱い",
        "第三者提供、業務委託、外部サービス利用（LINE、ホスティング、DB、AI等）",
        "共同利用がある場合の範囲・責任者",
        "保存期間と削除方針",
        "安全管理措置",
        "本人からの開示、訂正、利用停止、削除等の請求方法",
        "未成年者を受け付ける場合の方針",
        "Cookie・アクセスログ等の取扱い",
        "改定方法、施行日・版番号",
    ):
        add_checkbox(doc, item)
    add_response_box(doc, "添付ファイル名／法務確認者／確認日：", lines=3)

    add_section_heading(doc, "C-3. 会場・ブランド資料")
    for item in (
        "会場レイアウト図（受付、出入口、テーブル、避難動線を含む）",
        "50名分のテーブル・席一覧",
        "主催者ロゴ、イベントロゴ、正式表記ルール",
        "参加者へ案内する会場アクセス情報",
        "緊急時・災害時の会場ルール",
    ):
        add_checkbox(doc, item)
    add_response_box(doc, "添付ファイル名：", lines=0)

    add_section_heading(doc, "D. 暫定コンテンツの承認")
    add_paragraph(doc, "現在stagingに暫定設定があります。各項目について、承認または修正指示をご回答ください。", size=10.5)
    approvals = [
        ("感情カード8枚", "名称・説明・画像・表現の適切性"),
        ("Dream固定文3候補", "AI停止時にも表示する文章とブランド表現"),
        ("席案内の5問", "質問文・選択肢・『答えたくない』の表示・利用目的"),
        ("LINE通知文面", "本人連携、受付、結果等で送る案内文と送信者名"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2700, 4500, 2160])
    mark_header_row(table.rows[0])
    for idx, text in enumerate(("対象", "確認内容", "回答")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, ACCENT_DARK)
        set_cell_border(cell, ACCENT_DARK)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cell.paragraphs[0].add_run(text), size=9.5, bold=True, color=WHITE)
    for label, detail in approvals:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_border(cell)
        set_cell_shading(cells[0], GRAY)
        set_font(cells[0].paragraphs[0].add_run(label), size=9.5, bold=True)
        set_font(cells[1].paragraphs[0].add_run(detail), size=9.5)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cells[2].paragraphs[0].add_run("☐承認\n☐修正希望"), size=9)
    set_table_geometry(table, [2700, 4500, 2160])
    add_response_box(doc, "修正指示：", lines=5)

    add_section_heading(doc, "E. 当日運用情報")
    add_paragraph(doc, "当日の役割・障害対応を決めるため、以下をご回答ください。", size=10.5)
    operations = (
        "受付開始時刻・スタッフ集合時刻",
        "受付担当者数と使用予定スマートフォン台数",
        "席配置の最終承認者",
        "結果確定・LINE通知の最終承認者",
        "緊急時の責任者と連絡方法（連絡先詳細は別送可）",
        "会場Wi-Fiの有無、携帯回線の受信状況",
        "遅刻者、欠席者、途中退出者の扱い",
        "紙の受付表・席表を準備する担当者",
        "キャンセル受付期限と当日キャンセル方針",
        "撮影、取材、見学者がいる場合の運用",
    )
    for item in operations:
        add_checkbox(doc, item)
    add_response_box(doc, "回答・補足：", lines=7)

    add_section_heading(doc, "F. 提出・承認")
    add_note_box(doc, "提出前確認", "未確定項目には、確定予定日と担当者を必ず記載してください。規約本文と会場レイアウトが未提出の場合、申込受付開始および本番可否判定を完了できません。", caution=True)
    add_key_value_table(doc, [
        ("提出日", "　　　　年　　　月　　　日"),
        ("会社・団体名", "　　　　　　　　　　　　　　　　　　　　　　　　　"),
        ("回答責任者", "氏名：　　　　　　　　　　　　役職：　　　　　　　"),
        ("承認者", "氏名：　　　　　　　　　　　　役職：　　　　　　　"),
        ("連絡事項", "　　　　　　　　　　　　　　　　　　　　　　　　　\n　　　　　　　　　　　　　　　　　　　　　　　　　"),
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "SHIME 2026年8月8日イベント クライアント準備依頼書"
    doc.core_properties.subject = "本番設定・規約・運用資料の提出依頼"
    doc.core_properties.author = "SHIME Project"
    doc.core_properties.keywords = "SHIME, イベント, 本番準備, クライアント確認"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
